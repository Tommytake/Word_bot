import asyncio
import os
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement  # XML elementlar bilan ishlash uchun
from docx.text.paragraph import Paragraph

# Bot token
API_TOKEN = '8867882862:AAFTtN4teX4YQkEriUaSy7tKTm2JAxjrcMg' 
bot = Bot(token=API_TOKEN)
dp = Dispatcher(storage=MemoryStorage())

class DocForm(StatesGroup):
    matn1 = State()
    matn2 = State()
    photos = State()

def format_matn1(paragraph, text):
    paragraph.text = text
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

def fix_and_insert_matn2(doc, text):
    """
    Ushbu funksiya {matn2} markerini o'chirib, foydalanuvchi matnini 
    aniq yuqoridan pastga (A -> B -> C) ketma-ketlikda joylashtiradi.
    """
    lines = [line for line in text.split('\n') if line.strip()]
    if not lines:
        return

    for p in doc.paragraphs:
        if "{matn2}" in p.text:
            # 1. Shablondagi {matn2} qatorining o'ziga birinchi abzasni yozamiz
            p.text = lines[0]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(0.5)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
            
            # 2. Qolgan abzaslarni tartib bilan aynan shu paragrafdan KEYIN qo'shib boramiz
            current_p = p
            for line in lines[1:]:
                # Word hujjati uchun yangi paragraflar bloki (p) yaratamiz
                new_p_element = OxmlElement('w:p')
                # Uni joriy paragrafdan KEYIN joylashtiramiz
                current_p._element.addnext(new_p_element)
                
                # Elementni python-docx obyekti sifatida o'qiymiz
                inserted_p = Paragraph(new_p_element, doc)
                inserted_p.text = line
                inserted_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                inserted_p.paragraph_format.first_line_indent = Inches(0.5)
                for run in inserted_p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                
                # Keyingi qator buning tagiga tushishi uchun mo'ljalni yangilaymiz
                current_p = inserted_p
            break

@dp.message(Command("start"))
async def start(message: types.Message, state: FSMContext):
    await message.answer("1-matnni kiriting (Markazlashadi):")
    await state.set_state(DocForm.matn1)

@dp.message(DocForm.matn1)
async def get_matn1(message: types.Message, state: FSMContext):
    await state.update_data(matn1=message.text)
    await message.answer("2-matnni kiriting (Abzaslar bilan yozishingiz mumkin):")
    await state.set_state(DocForm.matn2)

@dp.message(DocForm.matn2)
async def get_matn2(message: types.Message, state: FSMContext):
    await state.update_data(matn2=message.text)
    await message.answer("Endi 2 ta rasmni yuboring:")
    await state.set_state(DocForm.photos)

@dp.message(DocForm.photos, F.photo)
async def get_photos(message: types.Message, state: FSMContext):
    data = await state.get_data()
    photos = data.get("photos", [])
    
    if len(photos) >= 2:
        return

    photo = message.photo[-1]
    file = await bot.get_file(photo.file_id)
    path = f"{message.from_user.id}_{len(photos)}.jpg"
    await bot.download_file(file.file_path, path)
    photos.append(path)
    await state.update_data(photos=photos)

    if len(photos) < 2:
        await message.answer(f"1-rasm qabul qilindi. Keyingi rasmni yuboring:")
        return

    await message.answer("Rasmlar qabul qilindi. Hujjat tayyorlanmoqda...")
    
    try:
        doc = Document("template.docx")
        
        # Matn 1 ni almashtirish
        for p in doc.paragraphs:
            if "{matn1}" in p.text:
                format_matn1(p, p.text.replace("{matn1}", data['matn1']))
                break
                
        # Matn 2 ni o'z o'rniga TO'G'RI TARTIBDA joylashtirish
        fix_and_insert_matn2(doc, data['matn2'])
        
        # Jadvalga rasmlarni joylash
        table = doc.add_table(rows=1, cols=2)
        for i, p_path in enumerate(photos):
            cell = table.cell(0, i)
            run = cell.paragraphs[0].add_run()
            run.add_picture(p_path, width=Inches(3.0))

        result_path = f"result_{message.from_user.id}.docx"
        doc.save(result_path)
        
        await message.answer_document(
            types.FSInputFile(result_path), 
            caption="Sizning hujjatingiz tayyor! 📄"
        )
    
    except Exception as e:
        await message.answer(f"Hujjat yaratishda xatolik yuz berdi: {str(e)}")
        
    finally:
        await state.clear()
        for p in photos:
            if os.path.exists(p): 
                os.remove(p)
        if 'result_path' in locals() and os.path.exists(result_path):
            os.remove(result_path)

async def main():
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())